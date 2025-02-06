from datetime import date
from Interface.CriaPerguntas import *
from VerificaçõesDeDados.VerificaçõesCliente import *
from Interface.MostraMensagem import *
from tinydb import Query
from FunçõesPrincipais.DevoluçãoEstoque import devolucao
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

Devolucao = Query()
Cliente = Query()

def criaRelatorio(interface, id):
    interface.withdraw()
    janela = CTkToplevel()
    janela.resizable(False, False)
    janela.geometry("520x330")
    janela.title("Relatórios")
    textoAviso = CTkLabel(janela, text="Clique em um dos botões para fazer a ação desejada!", font=("roboto", 20), text_color=corFonte)
    textoAviso.grid(padx=10, pady=13)

    janela.configure(fg_color=corInterfacePrincipal)
    frame = CTkFrame(janela, width=570, height=160, fg_color=corFundoFrame, border_color=corBordas, border_width=2)
    frame.grid(padx=25, pady=30)

    cliente = db.get(Cliente.id == int(id))

    informaçõesUsuario = {
        "nomeCliente": cliente['nome'],
        "cpfCnpjCliente": cliente['cpf'] if cliente['cpf'] != 'n/a' else cliente['cnpj'],
        "endereçoCliente": cliente['addresClient'],
        "endereçoObra": cliente['addresObra'],
        "numeroContrato": cliente['contractNumber']
    }

    def formata_texto(celula, texto):
        celula.text = texto
        for paragrafo in celula.paragraphs:
            for run in paragrafo.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(11)
                r = run._element
                r.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")

    def criaRelatorioContrato():
        try:
            relatorioContrato = Document('C:\\Users\\vinic\\OneDrive\\Documentos\\Área de Trabalho\\DocumentosClientes\\Modelos\\modeloRelatorioContrato.docx')

            tabelaLocatario = relatorioContrato.tables[0]

            for i, row in enumerate(tabelaLocatario.rows):
                for cell in row.cells:
                    if i > 0:
                        cell.text = cell.text.replace("nomeCliente", informaçõesUsuario["nomeCliente"])
                        cell.text = cell.text.replace("cpfCnpjCliente", informaçõesUsuario["cpfCnpjCliente"])
                        cell.text = cell.text.replace("endereçoCliente", informaçõesUsuario["endereçoCliente"])
                        cell.text = cell.text.replace("endereçoObra", informaçõesUsuario["endereçoObra"])
                        formata_texto(cell, cell.text)

            for paragrafo in relatorioContrato.paragraphs:
                for run in paragrafo.runs:
                    if "numeroContrato" in run.text:
                        run.text = run.text.replace("numeroContrato", informaçõesUsuario["numeroContrato"])

            tabela = relatorioContrato.tables[1]

            verificacaoItens = {
                "quantity": cliente["quantity"],
                "Andaime": cliente["quantityAndaime"],
                "Roldana": cliente["quantityRoldana"],
                "Plataforma": cliente["quantityPlataforma"],
                "Betoneira": cliente["quantityBetoneira"],
                "Regulador": cliente["quantityRegulador"]
            }

            chaveItens = [chave for chave, qtd in verificacaoItens.items() if qtd != 0]

            mapeamento_produtos = {
                "quantity": {"patrimonio": lambda c: "001" if c["sizeEscora"] == "3" else "009",
                            "descricao": lambda c: f"ESCORA {c['sizeEscora']}M",
                            "preco": lambda c: "300" if c["sizeEscora"] == "3" else "350",
                            "valor": "valueEscora"},
                
                "Andaime": {"patrimonio": lambda c: "003" if c["sizeAndaime"] == "1" else "002",
                            "descricao": lambda c: f"ANDAIME {c['sizeAndaime']}M",
                            "preco": lambda c: "200" if c["sizeAndaime"] == "1" else "220",
                            "valor": "valueAndaimes"},
                
                "Plataforma": {"patrimonio": lambda c: "005" if c["sizePlataforma"] == "1" else "004",
                            "descricao": lambda c: f"PLATAFORMA {c['sizePlataforma']}M",
                            "preco": lambda c: "180" if c["sizePlataforma"] == "1" else "200",
                            "valor": "valuePlataforma"},
                
                "Roldana": {"patrimonio": "006", "descricao": "ROLDANA", "preco": "180", "valor": "valueRoldana"},
                
                "Regulador": {"patrimonio": "007", "descricao": "REGULADOR", "preco": "100", "valor": "valueRegulador"},
                
                "Betoneira": {"patrimonio": "008", "descricao": "BETONEIRA", "preco": "5000", "valor": "valueBetoneira"}
            }

            for produto in chaveItens:
                novoProduto = tabela.add_row()
                config = mapeamento_produtos[produto]

                novoProduto.cells[0].text = config["patrimonio"](cliente) if callable(config["patrimonio"]) else config["patrimonio"]
                novoProduto.cells[1].text = config["descricao"](cliente) if callable(config["descricao"]) else config["descricao"]
                novoProduto.cells[2].text = verificacaoItens[produto]
                novoProduto.cells[3].text = config["preco"](cliente) if callable(config["preco"]) else config["preco"]
                novoProduto.cells[4].text = cliente[config["valor"]]

            linhaPagamento = tabela.add_row()
            linhaPagamento.cells[0].text = f"Total Geral: R${cliente['paymentAmount']}"
            linhaPagamento.cells[1].text = f"Desconto: R$"
            linhaPagamento.cells[2].text = f"Valor a pagar: R${cliente['paymentAmount']}"
            linhaPagamento.cells[3].text = f"Vencimento: {cliente['todayDate']} à {cliente['paymentDate']}"
            linhaPagamento._element.remove(linhaPagamento.cells[4]._element)

            relatorioContrato.save(r'C:\Users\vinic\OneDrive\Documentos\Área de Trabalho\DocumentosClientes\RelatoriosContratos\relatorioContrato{}.docx'.format(cliente['nome'].capitalize()))
            mostraMensagem("Ação bem-sucedida!", "O contrato foi criado com sucesso!")

        except Exception as e:
            mostraMensagem("Erro ao criar contrato", str(e))

    def criaRelatorioDevoluçõesCliente():
        try:
            relatorioContrato = Document('C:\\Users\\vinic\\OneDrive\\Documentos\\Área de Trabalho\\DocumentosClientes\\Modelos\\modeloDevoluçãoProdutos.docx')

            tabelaLocatario = relatorioContrato.tables[0]

            for i, row in enumerate(tabelaLocatario.rows):
                for cell in row.cells:
                    if i > 0:
                        cell.text = cell.text.replace("nomeCliente", informaçõesUsuario["nomeCliente"])
                        cell.text = cell.text.replace("cpfCnpjCliente", informaçõesUsuario["cpfCnpjCliente"])
                        cell.text = cell.text.replace("endereçoCliente", informaçõesUsuario["endereçoCliente"])
                        cell.text = cell.text.replace("endereçoObra", informaçõesUsuario["endereçoObra"])
                        formata_texto(cell, cell.text)
            
            for paragrafo in relatorioContrato.paragraphs:
                for run in paragrafo.runs:
                    if "nomeCliente" in run.text:
                        run.text = run.text.replace("nomeCliente", informaçõesUsuario["nomeCliente"])
                    if "DATA_____/____/_________" in run.text:
                        run.text = run.text.replace("DATA_____/____/_________", date.today().strftime('%d/%m/%Y'))
                    if "numeroContrato" in run.text:
                        run.text = run.text.replace("numeroContrato", informaçõesUsuario["numeroContrato"])

            mapeamento_devoluções = {
                "quantity": {"descricao": lambda c: f"ESCORA {c['sizeEscora']}M",
                            "patrimonio": lambda c: "001" if c["sizeEscora"] == "3" else "009",
                            "devolucao": "quantidadeDevolvida",
                            "restante": "quantidadeRestante",
                            "data": "dataDevolucao"},

                "quantityAndaime": {"descricao": lambda c: f"ANDAIME {c['sizeAndaime']}M",
                                    "patrimonio": lambda c: "003" if c["sizeAndaime"] == "1" else "002",
                                    "devolucao": "quantidadeDevolvida",
                                    "restante": "quantidadeRestante",
                                    "data": "dataDevolucao"},

                "quantityBetoneira": {"descricao": "BETONEIRA",
                                    "patrimonio": "008",
                                    "devolucao": "quantidadeDevolvida",
                                    "restante": "quantidadeRestante",
                                    "data": "dataDevolucao"},

                "quantityRoldana": {"descricao": "ROLDANA",
                                    "patrimonio": "006",
                                    "devolucao": "quantidadeDevolvida",
                                    "restante": "quantidadeRestante",
                                    "data": "dataDevolucao"},

                "quantityPlataforma": {"descricao": lambda c: f"PLATAFORMA {c['sizePlataforma']}M",
                                    "patrimonio": lambda c: "005" if c["sizePlataforma"] == "1" else "004",
                                    "devolucao": "quantidadeDevolvida",
                                    "restante": "quantidadeRestante",
                                    "data": "dataDevolucao"},
                
                "quantityRegulador": {"descricao": "REGULADOR",
                                    "patrimonio": "007",
                                    "devolucao": "quantidadeDevolvida",
                                    "restante": "quantidadeRestante",
                                    "data": "dataDevolucao"}
            }

            tabelaProdutos = relatorioContrato.tables[1]
            devoluções = devolucao.all()

            for devoluçõesClientes in devoluções:
                if devoluçõesClientes['id'] == cliente['id']:
                    novaDevolução = tabelaProdutos.add_row()
                    config = mapeamento_devoluções[devoluçõesClientes['tipoDevolucao']]

                    novaDevolução.cells[0].text = config["descricao"](cliente) if callable(config["descricao"]) else config["descricao"]
                    novaDevolução.cells[1].text = config["patrimonio"](cliente) if callable(config["patrimonio"]) else config["patrimonio"]
                    novaDevolução.cells[2].text = devoluçõesClientes[config["devolucao"]]
                    novaDevolução.cells[3].text = devoluçõesClientes[config["restante"]]
                    novaDevolução.cells[4].text = devoluçõesClientes[config["data"]]

            relatorioContrato.save(r'C:\Users\vinic\OneDrive\Documentos\Área de Trabalho\DocumentosClientes\RelatoriosDevoluções\relatorioDevolução{}.docx'.format(cliente['nome'].capitalize()))
            mostraMensagem("Ação bem-sucedida!", "O contrato de devoluções foi criado com sucesso!")

        except Exception as e:
            mostraMensagem("Erro ao criar contrato de devoluções", str(e))

    def criaReciboContrato():
        try:
            relatorioRecibo = Document('C:\\Users\\vinic\\OneDrive\\Documentos\\Área de Trabalho\\DocumentosClientes\\Modelos\\modeloRecibo.docx')

            tabelaLocatario = relatorioRecibo.tables[1]

            for i, row in enumerate(tabelaLocatario.rows):
                for cell in row.cells:
                    if i > 0:
                        cell.text = cell.text.replace("nomeCliente", informaçõesUsuario["nomeCliente"])
                        cell.text = cell.text.replace("cpfCnpjCliente", informaçõesUsuario["cpfCnpjCliente"])
                        cell.text = cell.text.replace("endereçoCliente", informaçõesUsuario["endereçoCliente"])
                        cell.text = cell.text.replace("endereçoObra", informaçõesUsuario["endereçoObra"])
                        formata_texto(cell, cell.text)
            
            for paragrafo in relatorioRecibo.paragraphs:
                paragrafo.text = paragrafo.text.replace("dataAtual", date.today().strftime('%d/%m/%Y'))

            tabelaDiscriminação = relatorioRecibo.tables[2]

            verificacaoItens = {
                "quantity": cliente["quantity"],
                "Andaime": cliente["quantityAndaime"],
                "Roldana": cliente["quantityRoldana"],
                "Plataforma": cliente["quantityPlataforma"],
                "Betoneira": cliente["quantityBetoneira"],
                "Regulador": cliente["quantityRegulador"]
            }

            chaveItens = [chave for chave, qtd in verificacaoItens.items() if qtd != 0]

            mapeamento_produtos = {
                "quantity": {
                    "descricao": lambda c: f"ESCORA {c['sizeEscora']}M",
                    "periodoContrato": lambda c: f"{c['todayDate']} à {c['paymentDate']}",
                    "valor": "valueEscora"
                },
                
                "Andaime": {
                    "descricao": lambda c: f"ANDAIME {c['sizeAndaime']}M",
                    "periodoContrato": lambda c: f"{c['todayDate']} à {c['paymentDate']}",
                    "valor": "valueAndaimes"
                },
                
                "Plataforma": {
                    "descricao": lambda c: f"PLATAFORMA {c['sizePlataforma']}M",
                    "periodoContrato": lambda c: f"{c['todayDate']} à {c['paymentDate']}",
                    "valor": "valuePlataforma"
                },
                
                "Roldana": {
                    "descricao": "ROLDANA", 
                    "periodoContrato": lambda c: f"{c['todayDate']} à {c['paymentDate']}", 
                    "valor": "valueRoldana"
                },
                
                "Regulador": {
                    "descricao": "REGULADOR", 
                    "periodoContrato": lambda c: f"{c['todayDate']} à {c['paymentDate']}", 
                    "valor": "valueRegulador"
                },
                
                "Betoneira": {
                    "descricao": "BETONEIRA", 
                    "periodoContrato": lambda c: f"{c['todayDate']} à {c['paymentDate']}", 
                    "valor": "valueBetoneira"
                }
            }

            primeira_iteracao = True

            for produto in chaveItens:
                novoProduto = tabelaDiscriminação.add_row()
                config = mapeamento_produtos[produto]

                novoProduto.cells[0].text = config["descricao"](cliente) if callable(config["descricao"]) else config["descricao"]
                novoProduto.cells[1].text = str(verificacaoItens[produto])
                novoProduto.cells[2].text = f"R$ {str(cliente[config["valor"]])}"
                novoProduto.cells[3].text = config["periodoContrato"](cliente) if callable(config["periodoContrato"]) else config["periodoContrato"]

                if primeira_iteracao:
                    novoProduto.cells[4].text = f"R$ {str(cliente["paymentAmount"])}"
                    primeira_iteracao = False
                else:
                    novoProduto._element.remove(novoProduto.cells[4]._element)

            linha_legislacao = tabelaDiscriminação.add_row()
            linha_legislacao.cells[0].merge(linha_legislacao.cells[-1])

            texto_legislacao = """RECIBO EMITIDO CONFORME A LEGISLAÇÃO EM VIGOR
            LC 123/2006, ARTIGO 18
            § 1 ISS PREVISTO NESTE ANEXO.
            AS ATIVIDADES DE LOCAÇÃO DE BENS MÓVEL SERÃO TRIBUTADAS NA FORMA DO ANEXO III DESTA LEI COMPLEMENTAR, DEDUZINDO-SE DA ALÍQUOTA O PERCENTUAL CORRESPONDENTE AO ISS PREVISTO NESTE ANEXO.
            A LOCAÇÃO DE BENS IMÓVEIS OU MÓVEL NÃO CONSTITUI UMA PRESTAÇÃO DE SERVIÇOS, MAS DISPONIBILIZAÇÃO DE UM BEM, SEJA ELE IMÓVEL OU MÓVEL PARA UTILIZAÇÃO DO LOCATÁRIO SEM A PRESTAÇÃO DE UM SERVIÇO.
            NOS ART. 18 A 20 DA CITADA LC. DOCUMENTO FISCAL: RECIBO. O RECIBO CONSTITUI DOCUMENTO HÁBIL PARA A COMPROVAÇÃO DE RECEITA AUFERIDA PELA PESSOA JURÍDICA OPTANTE PELO SIMPLES QUE NÃO SEJA OBRIGADA À EMISSÃO DE NOTA FISCAL OU À UTILIZAÇÃO DE EQUIPAMENTO EMISSOR DE CUPOM FISCAL - ECF"""

            linha_legislacao.cells[0].text = texto_legislacao

            paragrafo = linha_legislacao.cells[0].paragraphs[0]

            run = paragrafo.runs[0]
            run.font.name = "Calibri"
            run.font.size = Pt(8)

            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), "Calibri")
            rFonts.set(qn("w:hAnsi"), "Calibri")
            rFonts.set(qn("w:eastAsia"), "Calibri")
            rFonts.set(qn("w:cs"), "Calibri")
            rPr.append(rFonts)

            caminho_arquivo = r'C:\Users\vinic\OneDrive\Documentos\Área de Trabalho\DocumentosClientes\RecibosClientes\recibo{}.docx'.format(cliente['nome'].capitalize())
            relatorioRecibo.save(caminho_arquivo)
            mostraMensagem("Ação bem-sucedida!", "O recibo foi criado com sucesso!")

        except Exception as e:
            mostraMensagem("Erro ao criar recibo", str(e))
            print(str(e))


    criaBotao(frame, "Equipamentos do Contrato", criaRelatorioContrato, 0, 0)
    criaBotao(frame, "Devolução de Equipamentos", criaRelatorioDevoluçõesCliente, 0, 1)
    criaBotao(frame, "Recibo", criaReciboContrato, 1, 0)

    def fechaJanela():
        janela.destroy()
        interface.deiconify()
    criaBotao(frame, "Voltar para o Inicio", fechaJanela, 1, 1)
    
    janela.protocol("WM_DELETE_WINDOW", fechaJanela)